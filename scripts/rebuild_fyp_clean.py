#!/usr/bin/env python3
"""
Rebuild KAEL FYP as a clean Times New Roman PDF (+ DOCX) from the original
Canva PDF text/images — NOT from the broken pdf2docx conversion.
"""

from __future__ import annotations

import re
from pathlib import Path
from collections import defaultdict

import pymupdf
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    Table,
    TableStyle,
)
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Cm as DocxCm, Pt as DocxPt, Inches as DocxInches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/fatee/Downloads")
SRC_PDF = ROOT / "KAEL_Report.pdf"
OUT_PDF = ROOT / "KAEL_FYP_Formatted.pdf"
OUT_DOCX = ROOT / "KAEL_FYP_Formatted.docx"
ASSETS = Path("/Users/fatee/Desktop/P R O J E C T/MIRVÉ-PORTFOLIO/.fyp_assets")

TIMES = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
TIMES_B = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"
TIMES_I = "/System/Library/Fonts/Supplemental/Times New Roman Italic.ttf"
TIMES_BI = "/System/Library/Fonts/Supplemental/Times New Roman Bold Italic.ttf"

# Margins per FYP spec
LEFT = 3.81 * cm
RIGHT = 2.54 * cm
TOP = 2.54 * cm
BOTTOM = 2.54 * cm

CHAPTERS = [
    (1, "INTRODUCTION"),
    (2, "EXISTING SYSTEM"),
    (3, "SOFTWARE PROCESS MODEL"),
    (4, "PROPOSED SYSTEM"),
    (5, "SYSTEM DESIGN"),
    (6, "DEVELOPMENT"),
    (7, "SOFTWARE TESTING"),
    (8, "IMPLEMENTATION AND TRAINING"),
]

TOC_ENTRIES = [
    (0, "ACKNOWLEDGEMENT", "ii"),
    (0, "ABSTRACT", "iii"),
    (0, "TABLE OF CONTENTS", "v"),
    (0, "LIST OF FIGURES", "viii"),
    (0, "LIST OF TABLES", "x"),
    (0, "CHAPTER 1: INTRODUCTION", "1"),
    (1, "1.1 INTRODUCTION", "1"),
    (1, "1.2 Domain", "1"),
    (2, "Subdomain", "2"),
    (1, "1.3 Problem Statement", "2"),
    (1, "1.4 Motivation", "2"),
    (1, "1.5 Definition of Terms", "3"),
    (1, "1.6 Workflow", "3"),
    (1, "1.7 Goal of Project", "4"),
    (0, "CHAPTER 2: EXISTING SYSTEM", "5"),
    (1, "2.1 Relative Work", "5"),
    (1, "2.2 Existing Websites and their features", "5"),
    (1, "2.3 Novelty", "6"),
    (1, "2.4 Solution Available in KAEL", "6"),
    (1, "2.5 System Requirement Specification", "6"),
    (2, "2.5.1 User Interfaces", "6"),
    (2, "2.5.2 Hardware Interfaces", "6"),
    (2, "2.5.3 Software Interfaces", "7"),
    (1, "2.6 Functional Requirements", "7"),
    (2, "2.6.1 Use Case Diagram", "9"),
    (2, "2.6.2 Requirement Shells", "10"),
    (1, "2.7 Non Functional Requirements", "16"),
    (2, "2.7.1 Shell Requiremnets", "16"),
    (0, "CHAPTER 3: SOFTWARE PROCESS MODEL", "19"),
    (1, "3.1 Software Process Model", "19"),
    (1, "3.2 Proposed Model", "20"),
    (2, "3.2.1 Why Agile Scrum was chosen", "20"),
    (1, "3.3 Implementation of Agile Scrum", "21"),
    (0, "CHAPTER 4: PROPOSED SYSTEM", "23"),
    (1, "4.1 Need for the proposed system", "23"),
    (1, "4.2 Features of Proposed System", "23"),
    (2, "4.2.1 User Data Collection", "23"),
    (2, "4.2.2 AI Processing Pipeline", "24"),
    (2, "4.2.3 Automated generation of case studies", "24"),
    (2, "4.2.4 Portfolio Builder", "24"),
    (2, "4.2.5 ATS-Friendly CV Builder", "24"),
    (1, "4.3 Business Context", "24"),
    (2, "4.3.1 Advantages in Business", "25"),
    (2, "4.3.2 Stakeholder's Profiles", "26"),
    (2, "4.3.3 Project Priorities", "27"),
    (1, "4.4 Operating Environment", "27"),
    (2, "4.4.1 Client-Side Requirements", "27"),
    (2, "4.4.2 Server-Side / Backend Requirements", "28"),
    (2, "4.4.3 Security and Compliance Considerations", "28"),
    (0, "CHAPTER 5: SYSTEM DESIGN", "29"),
    (1, "5.1 Brief Description", "29"),
    (1, "5.2 UI/UX Design", "29"),
    (2, "1. Landing Page", "30"),
    (2, "2. Sign Up", "31"),
    (2, "3. Onboarding & Visual Identity Curation", "31"),
    (2, "4. Profile Creation", "33"),
    (2, "5. Case Study Generator", "35"),
    (2, "6. CV Generator", "36"),
    (2, "7. Portfolio Generator", "38"),
    (1, "5.3 ER Diagram for KAEL", "40"),
    (1, "5.4 DFD Diagram", "40"),
    (2, "Level 0", "40"),
    (2, "Level 1", "41"),
    (2, "Level 2", "42"),
    (1, "5.5 Activity Diagram KAEL", "43"),
    (1, "5.6 Class Diagram", "47"),
    (1, "5.7 Software Architectural Design", "47"),
    (1, "5.8 Sequence Diagram", "48"),
    (0, "CHAPTER 6: DEVELOPMENT", "49"),
    (1, "6.1 Selection of Tools & Technology", "49"),
    (2, "6.1.1 Hardware", "49"),
    (2, "6.1.2 Software", "49"),
    (0, "CHAPTER 7: SOFTWARE TESTING", "51"),
    (1, "7.1 Software Testing", "51"),
    (1, "7.2 Testing Process", "51"),
    (1, "7.3 Black Box Testing", "52"),
    (1, "7.4 Test Cases", "53"),
    (0, "CHAPTER 8: IMPLEMENTATION AND TRAINING", "56"),
    (1, "8.1 Implementation", "56"),
    (1, "8.2 Project Overview (Key Features only)", "57"),
    (2, "8.2.1 Expected Outcome", "57"),
    (1, "8.3 Training", "58"),
    (2, "8.3.1 Required Skillset", "58"),
    (2, "8.3.2 User Training Protocol", "58"),
    (0, "CONCLUSION", "60"),
    (0, "REFERENCES", "61"),
    (0, "FINAL YEAR PROJECT UNDERTAKING FORM", "62"),
]

LOF = [
    ("Figure 1.1 KAEL Gantt Chart", "4"),
    ("Figure 2.1 Use case diagram - Student", "9"),
    ("Figure 2.2 Use case diagram - Recruiter", "9"),
    ("Figure 2.3 Use case diagram - Academic Authority", "10"),
    ("Figure 3.1 Agile scrum cycle adopted for KAEL development process", "20"),
    ("Figure 5.1 KAEL landing page", "30"),
    ("Figure 5.2 Sign-Up Screen", "31"),
    ("Figure 5.3 Creative-Academic Duality Slider", "32"),
    ("Figure 5.4 Profile Photo Upload", "32"),
    ("Figure 5.5 Educational Background Form", "32"),
    ("Figure 5.6 Soft Skills & Visual Identity", "33"),
    ("Figure 5.7 Soft Skills Selection", "33"),
    ("Figure 5.8 Profile Editor (Dark Theme)", "34"),
    ("Figure 5.9 Profile Editor (Light Theme)", "34"),
    ("Figure 5.10 Case Study Creation Interface", "35"),
    ("Figure 5.11 Generated Case Study Preview", "36"),
    ("Figure 5.12 CV Builder Form", "36"),
    ("Figure 5.13 Job Description Input", "37"),
    ("Figure 5.14 Job Match Analysis", "37"),
    ("Figure 5.15 Final ATS CV Preview", "37"),
    ("Figure 5.16 Portfolio Template Selection", "38"),
    ("Figure 5.17 Portfolio Live Preview", "39"),
    ("Figure 5.18 Portfolio Colour Customization", "39"),
    ("Figure 5.19 Customized Portfolio", "39"),
    ("Figure 5.20 ER Diagram of KAEL", "40"),
    ("Figure 5.21 DFD Level 0 - Context Diagram", "40"),
    ("Figure 5.22 DFD Level 1 - Onboarding, Case Study, Portfolio & Job Match Processes", "41"),
    ("Figure 5.23 DFD Level 2 - Detailed Process Decomposition", "42"),
    ("Figure 5.24 User Onboarding (Activity Diagram)", "43"),
    ("Figure 5.25 Case Study Creation (Activity Diagram)", "44"),
    ("Figure 5.26 Portfolio Generation & Publishing", "45"),
    ("Figure 5.27 Job Match & Tailored CV Generation", "46"),
    ("Figure 5.28 Class Diagram of KAEL", "47"),
    ("Figure 5.29 Software Architectural Design", "47"),
    ("Figure 5.30 Sequence Diagram of KAEL", "48"),
]

LOT = [
    ("Table 1.1 Explanation of terms used in documentation", "3"),
    ("Table 2.1 Comparison with existing related web applications", "5"),
    ("Table 2.2 Software interfaces used", "7"),
    ("Table 4.1 Stakeholder's profiles", "26"),
    ("Table 4.2 Project Priorities of KAEL", "27"),
    ("Table 7.1 Test cases for KAEL", "53"),
]

FUNC_BULLETS = [
    "Profile Management: Gives an interactive canvas for users to organize their personal information, education, skills and professional objectives. This is the basic metadata which is used as the main context in AI generation and layout engines downstream.",
    "Project Upload: Consumes diverse engineering files such as heavy PDF documents, documentation reports, source code files, screenshots and demonstration videos. These unstructured assets are fed directly into separate non-corruptible Firebase Storage buckets.",
    "AI Project Analysis: Extracts information in a multimodal way from uploaded documents to capture underlying project targets, tools and technical competencies. It removes the need for manual tracking, and automatically extracts accurate performance data directly from raw files.",
    "Automatic Case Study Generation: Applies rich contextual cues to transform clumps of academic project information into well-formulated domain-specific case studies tailored to the conventions of that specific academic or industry domain.",
    "AI Portfolio Generation: Cleanly groups student assets by technical domains and automatically designs and structures accessible web portfolio dashboards.",
    "ATS CV Generator: Exports extracted skills and project histories to cleanly structured CVs with one column, dynamically. Removes parsing errors, optimizing the final document for ATS systems.",
    "Academic Document Upload: Serves as a private cloud vault for the safe storage of everyday lecture notes, assignments, research papers, and certificates.",
    "Soft Skill Analysis: Compares non-technical performance attributes, such as rate of change in assets, document structure decisions, and quality of the writing over time.",
    "Academic Performance Matrix: Gathers and presents department grades, skill points, project feedback milestones in a neat visual dashboard for analytics.",
    "Portfolio & CV Export: Creates unauthenticated public sharing web subdomains, and enables quick local downloads of printable PDF resumes.",
    "AI Job Match Analyzer: Compares a candidate's current engineering description to job posting requirements through a real-time semantic comparison of keywords.",
]

NFR_BULLETS = [
    "Security & Data Protection",
    "Performance & Responsiveness",
    "Scalability & Availability",
    "Usability & Accessibility",
]


def register_fonts():
    pdfmetrics.registerFont(TTFont("Times", TIMES))
    pdfmetrics.registerFont(TTFont("Times-Bold", TIMES_B))
    pdfmetrics.registerFont(TTFont("Times-Italic", TIMES_I))
    pdfmetrics.registerFont(TTFont("Times-BoldItalic", TIMES_BI))


def make_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            fontName="Times-Bold",
            fontSize=28,
            alignment=TA_CENTER,
            spaceAfter=8,
            leading=34,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverSub",
            fontName="Times-Italic",
            fontSize=14,
            alignment=TA_CENTER,
            spaceAfter=20,
            leading=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverBody",
            fontName="Times",
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=6,
            leading=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ChapPage",
            fontName="Times-Bold",
            fontSize=22,
            alignment=TA_CENTER,
            spaceAfter=12,
            leading=28,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ChapPageName",
            fontName="Times-Bold",
            fontSize=18,
            alignment=TA_CENTER,
            leading=24,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HChapter",
            fontName="Times-Bold",
            fontSize=16,
            alignment=TA_LEFT,
            spaceBefore=6,
            spaceAfter=12,
            leading=20,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HSection",
            fontName="Times-Bold",
            fontSize=14,
            alignment=TA_LEFT,
            spaceBefore=12,
            spaceAfter=6,
            leading=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HSub",
            fontName="Times-Bold",
            fontSize=12,
            alignment=TA_LEFT,
            spaceBefore=8,
            spaceAfter=4,
            leading=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyJust",
            fontName="Times",
            fontSize=12,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
            leading=16,
            firstLineIndent=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Caption",
            fontName="Times-Italic",
            fontSize=11,
            alignment=TA_CENTER,
            spaceBefore=4,
            spaceAfter=12,
            leading=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCTitle",
            fontName="Times-Bold",
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=18,
            leading=20,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCEntry",
            fontName="Times",
            fontSize=12,
            alignment=TA_LEFT,
            spaceAfter=3,
            leading=15,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletBody",
            fontName="Times",
            fontSize=12,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
            leading=16,
            leftIndent=15,
            bulletIndent=0,
        )
    )
    return styles


def escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def is_page_num(line: str) -> bool:
    return bool(re.fullmatch(r"\d+|i{1,3}|iv|v|vi{0,3}|ix|x|xi{0,2}", line.strip(), re.I))


def is_heading(line: str) -> tuple[str | None, str]:
    t = line.strip()
    if not t:
        return None, t
    if re.match(r"^CHAPTER\s+\d+:", t, re.I):
        return "chapter", t
    if t.upper() in {
        "ACKNOWLEDGEMENT",
        "ABSTRACT",
        "CONCLUSION",
        "REFERENCES",
        "INTRODUCTION",
        "SUBDOMAIN",
    }:
        return "section", t
    if re.match(r"^\d+\.\d+(\.\d+)?\s+\S", t):
        # 1.1 / 2.5.1 style
        if re.match(r"^\d+\.\d+\.\d+", t):
            return "sub", t
        return "section", t
    if re.match(r"^\d+\.\s+[A-Z]", t):  # 1. Landing Page
        return "sub", t
    if t.startswith("Figure ") or t.startswith("Table "):
        return "caption", t
    if t in {
        "Client-Side Requirements",
        "Server-Side / Backend Requirements",
        "Security and Compliance Considerations",
        "Expected Outcome",
        "Required Skillset",
        "User Training Protocol",
        "Level 0",
        "Level 1",
        "Level 2",
    }:
        return "sub", t
    return None, t


def load_pages_text() -> dict[int, str]:
    doc = pymupdf.open(SRC_PDF)
    pages = {}
    for i, page in enumerate(doc):
        pages[i + 1] = page.get_text("text")
    doc.close()
    return pages


def largest_image_on_page(page_no: int) -> Path | None:
    files = sorted(ASSETS.glob(f"p{page_no:02d}_*.png"))
    if not files:
        # try without zero pad variants already zero-padded
        files = sorted(ASSETS.glob(f"p{page_no}_*.png"))
    if not files:
        return None
    files = sorted(files, key=lambda p: p.stat().st_size, reverse=True)
    # skip tiny icons
    for f in files:
        if f.stat().st_size > 8000:
            return f
    return files[0] if files else None


def clean_paragraphs(raw: str) -> list[tuple[str, str]]:
    """Return list of (kind, text) where kind in body|section|sub|chapter|caption|skip."""
    lines = [ln.strip() for ln in raw.splitlines()]
    out: list[tuple[str, str]] = []
    buf: list[str] = []

    def flush():
        nonlocal buf
        if buf:
            text = " ".join(buf)
            text = re.sub(r"\s+", " ", text).strip()
            if text and not is_page_num(text):
                out.append(("body", text))
            buf = []

    for ln in lines:
        if not ln or is_page_num(ln):
            flush()
            continue
        kind, _ = is_heading(ln)
        if kind:
            flush()
            out.append((kind, ln))
        else:
            buf.append(ln)
    flush()
    return out


def renumber_chapter1(items: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """Ensure Ch1 numbering: 1.1 INTRODUCTION, 1.2 Domain, ... 1.7 Goal."""
    mapping = {
        "INTRODUCTION": "1.1 INTRODUCTION",
        "1.1 Domain": "1.2 Domain",
        "1.2 Domain": "1.2 Domain",
        "1.2 Problem Statement": "1.3 Problem Statement",
        "1.3 Problem Statement": "1.3 Problem Statement",
        "1.3 Motivation": "1.4 Motivation",
        "1.4 Motivation": "1.4 Motivation",
        "1.4 Definition of Terms": "1.5 Definition of Terms",
        "1.5 Definition of Terms": "1.5 Definition of Terms",
        "1.5 Workflow": "1.6 Workflow",
        "1.6 Workflow": "1.6 Workflow",
        "1.6 Goal of Project": "1.7 Goal of Project",
        "1.7 Goal of Project": "1.7 Goal of Project",
    }
    result = []
    for kind, text in items:
        if kind in {"section", "sub"} and text in mapping:
            result.append((kind if not mapping[text].startswith("1.") else "section", mapping[text]))
        elif kind == "section" and text.upper() == "INTRODUCTION":
            result.append(("section", "1.1 INTRODUCTION"))
        elif kind == "section" and text == "Domain":
            result.append(("section", "1.2 Domain"))
        else:
            result.append((kind, text))
    return result


def fix_unnumbered(items: list[tuple[str, str]]) -> list[tuple[str, str]]:
    repl = {
        "Client-Side Requirements": "4.4.1 Client-Side Requirements",
        "Server-Side / Backend Requirements": "4.4.2 Server-Side / Backend Requirements",
        "Security and Compliance Considerations": "4.4.3 Security and Compliance Considerations",
        "Expected Outcome": "8.2.1 Expected Outcome",
        "Required Skillset": "8.3.1 Required Skillset",
        "User Training Protocol": "8.3.2 User Training Protocol",
    }
    out = []
    for kind, text in items:
        if text in repl:
            out.append(("sub", repl[text]))
        else:
            out.append((kind, text))
    return out


def make_image(path: Path, max_width=14 * cm, max_height=10 * cm):
    img = Image(str(path))
    img.hAlign = "CENTER"
    aspect = img.imageHeight / float(img.imageWidth)
    w = max_width
    h = w * aspect
    if h > max_height:
        h = max_height
        w = h / aspect
    img.drawWidth = w
    img.drawHeight = h
    return img


def toc_table(entries, styles):
    """Build TOC-like table with dotted leaders."""
    data = []
    for level, title, page in entries:
        indent = "&nbsp;" * (level * 4)
        left = Paragraph(f"{indent}{escape(title)}", styles["TOCEntry"])
        right = Paragraph(escape(str(page)), styles["TOCEntry"])
        data.append([left, right])
    t = Table(data, colWidths=[14.5 * cm, 1.5 * cm])
    t.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    return t


def build_pdf(pages_text: dict[int, str]):
    register_fonts()
    styles = make_styles()
    story = []

    # ---- COVER ----
    cover_img = largest_image_on_page(1)
    story.append(Spacer(1, 1.5 * cm))
    if cover_img:
        try:
            story.append(make_image(cover_img, max_width=4 * cm, max_height=4 * cm))
        except Exception:
            pass
    story.append(Spacer(1, 1.2 * cm))
    story.append(Paragraph("KAEL", styles["CoverTitle"]))
    story.append(Paragraph("Desktop Application", styles["CoverSub"]))
    story.append(Spacer(1, 0.8 * cm))
    story.append(
        Paragraph(
            "A project submitted in partial fulfillment of the requirements<br/>"
            "for the degree of<br/><b>Bachelor of Science in Computer Science</b>",
            styles["CoverBody"],
        )
    )
    story.append(Spacer(1, 0.8 * cm))
    story.append(Paragraph("Hoorulain Fatima (2k22-BSCS-207)", styles["CoverBody"]))
    story.append(Paragraph("Fatima Usman (2k22-BSCS-225)", styles["CoverBody"]))
    story.append(Paragraph("Abeeha Fatima (2k22-BSCS-212)", styles["CoverBody"]))
    story.append(Paragraph("Khadija Sajid (2k22-BSCS-256)", styles["CoverBody"]))
    story.append(Spacer(1, 1.0 * cm))
    story.append(Paragraph("Supervised By", styles["CoverBody"]))
    story.append(Paragraph("<b>Sir Syed Umair Asghar</b>", styles["CoverBody"]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Head of Department", styles["CoverBody"]))
    story.append(Paragraph("<b>Dr. Naeem Aslam</b>", styles["CoverBody"]))
    story.append(Spacer(1, 1.0 * cm))
    story.append(Paragraph("DEPARTMENT OF COMPUTER SCIENCE", styles["CoverBody"]))
    story.append(
        Paragraph(
            "NFC Institute of Engineering &amp; Technology, Multan",
            styles["CoverBody"],
        )
    )
    story.append(PageBreak())

    # ---- ACKNOWLEDGEMENT ----
    story.append(Paragraph("ACKNOWLEDGEMENT", styles["TOCTitle"]))
    ack = pages_text.get(2, "")
    for kind, text in clean_paragraphs(ack):
        if kind == "body" and "ACKNOWLEDGEMENT" not in text and text.lower() != "ii":
            story.append(Paragraph(escape(text), styles["BodyJust"]))
    story.append(PageBreak())

    # ---- ABSTRACT ----
    story.append(Paragraph("ABSTRACT", styles["TOCTitle"]))
    ab = pages_text.get(3, "")
    for kind, text in clean_paragraphs(ab):
        if kind == "body" and text.upper() != "ABSTRACT" and text.lower() != "iii":
            story.append(Paragraph(escape(text), styles["BodyJust"]))
    story.append(PageBreak())

    # ---- TOC / LOF / LOT ----
    story.append(Paragraph("TABLE OF CONTENTS", styles["TOCTitle"]))
    story.append(toc_table(TOC_ENTRIES, styles))
    story.append(PageBreak())

    story.append(Paragraph("LIST OF FIGURES", styles["TOCTitle"]))
    story.append(toc_table([(0, a, b) for a, b in LOF], styles))
    story.append(PageBreak())

    story.append(Paragraph("LIST OF TABLES", styles["TOCTitle"]))
    story.append(toc_table([(0, a, b) for a, b in LOT], styles))
    story.append(PageBreak())

    # ---- BODY: process original pages 11–72 with chapter openers ----
    # Pre-group by chapter using page ranges from original document
    chapter_ranges = {
        1: range(11, 15),
        2: range(15, 29),
        3: range(29, 33),
        4: range(33, 39),
        5: range(39, 59),
        6: range(59, 61),
        7: range(61, 66),
        8: range(66, 70),
    }
    extras = {
        "CONCLUSION": range(70, 71),
        "REFERENCES": range(71, 72),
        "FORM": range(72, 73),
    }

    inserted_chapters = set()
    skip_until_after_func_list = False

    def add_chapter_opener(num: int, name: str):
        story.append(Spacer(1, 8 * cm))
        story.append(Paragraph(f"CHAPTER {num}", styles["ChapPage"]))
        story.append(Paragraph(name, styles["ChapPageName"]))
        story.append(PageBreak())

    # Walk pages in order
    for page_no in range(11, 73):
        # Chapter openers
        for num, name in CHAPTERS:
            start = min(chapter_ranges[num])
            if page_no == start and num not in inserted_chapters:
                add_chapter_opener(num, name)
                inserted_chapters.add(num)

        if page_no == 70:
            story.append(Spacer(1, 8 * cm))
            story.append(Paragraph("CONCLUSION", styles["ChapPage"]))
            story.append(PageBreak())
        if page_no == 71:
            story.append(Spacer(1, 8 * cm))
            story.append(Paragraph("REFERENCES", styles["ChapPage"]))
            story.append(PageBreak())

        raw = pages_text.get(page_no, "")
        items = clean_paragraphs(raw)
        if page_no in range(11, 15):
            items = renumber_chapter1(items)
        items = fix_unnumbered(items)

        # Special: Chapter 3 — put Figure 3.1 early under 3.2
        # Special: 2.6 bullets — replace feature dump with clean bullets once
        i = 0
        while i < len(items):
            kind, text = items[i]

            # Skip original broken TOC-like remnants / roman markers
            if text.lower() in {"ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x"}:
                i += 1
                continue

            if kind == "chapter":
                # Already have opener pages; still show chapter heading once on content page
                story.append(Paragraph(escape(text), styles["HChapter"]))
            elif kind == "section":
                # Inject clean 2.6 bullet list when we hit Functional Requirements
                if text.startswith("2.6 Functional Requirements"):
                    story.append(Paragraph(escape(text), styles["HSection"]))
                    for b in FUNC_BULLETS:
                        story.append(Paragraph(f"• {escape(b)}", styles["BulletBody"]))
                    # skip following body lines until 2.6.1 / 2.7
                    i += 1
                    while i < len(items):
                        k2, t2 = items[i]
                        if k2 in {"section", "sub"} and (
                            t2.startswith("2.6.1") or t2.startswith("2.7")
                        ):
                            break
                        if k2 in {"section", "sub", "chapter", "caption"}:
                            break
                        i += 1
                    continue
                if text.startswith("2.7 Non Functional"):
                    story.append(Paragraph(escape(text), styles["HSection"]))
                    for b in NFR_BULLETS:
                        story.append(Paragraph(f"• {escape(b)}", styles["BulletBody"]))
                    i += 1
                    continue
                story.append(Paragraph(escape(text), styles["HSection"]))
            elif kind == "sub":
                # Fix 7.4 heading fragments
                if text.startswith("7.4"):
                    story.append(Paragraph("7.4 Test Cases", styles["HSub"]))
                else:
                    story.append(Paragraph(escape(text), styles["HSub"]))
            elif kind == "caption":
                # Place image above caption when available
                img_path = largest_image_on_page(page_no)
                # For figure 3.1 prefer page 30 image
                if text.startswith("Figure 3.1"):
                    img_path = largest_image_on_page(30) or img_path
                if img_path and text.startswith("Figure"):
                    try:
                        story.append(Spacer(1, 6))
                        story.append(make_image(img_path))
                    except Exception:
                        pass
                story.append(Paragraph(escape(text), styles["Caption"]))
            else:
                # body
                # skip lines that are clearly UI chrome / gantt headers noise if very short keywords only
                if len(text) < 3:
                    i += 1
                    continue
                story.append(Paragraph(escape(text), styles["BodyJust"]))
            i += 1

        # After page content, if page has a dominant figure not captioned in text, still add image for UI pages
        # (handled via captions mostly)

    # Build
    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        leftMargin=LEFT,
        rightMargin=RIGHT,
        topMargin=TOP,
        bottomMargin=BOTTOM,
        title="KAEL Final Year Project Report",
        author="KAEL Team",
    )

    def add_page_number(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Times", 11)
        page = canvas.getPageNumber()
        canvas.drawCentredString(A4[0] / 2, 1.2 * cm, str(page))
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    print("Wrote PDF", OUT_PDF)


def build_docx_from_pdf_mirror():
    """Create a clean DOCX sibling with same structure/rules for editing."""
    register_fonts()  # not needed for docx but ok
    pages_text = load_pages_text()
    d = Document()

    # margins
    for sec in d.sections:
        sec.top_margin = DocxCm(2.54)
        sec.bottom_margin = DocxCm(2.54)
        sec.right_margin = DocxCm(2.54)
        sec.left_margin = DocxCm(3.81)
        sec.page_width = DocxCm(21.0)
        sec.page_height = DocxCm(29.7)

    style = d.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = DocxPt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(a), "Times New Roman")

    def add_center(text, size=16, bold=True):
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = DocxPt(size)
        run.bold = bold
        return p

    def add_just(text, size=12, bold=False):
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = DocxPt(size)
        run.bold = bold
        return p

    def add_heading_left(text, size=14):
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = DocxPt(size)
        run.bold = True
        return p

    def page_break():
        d.add_page_break()

    # Cover
    add_center("KAEL", 28, True)
    add_center("Desktop Application", 14, False)
    add_center(
        "A project submitted in partial fulfillment of the requirements for the degree of Bachelor of Science in Computer Science",
        12,
        False,
    )
    for name in [
        "Hoorulain Fatima (2k22-BSCS-207)",
        "Fatima Usman (2k22-BSCS-225)",
        "Abeeha Fatima (2k22-BSCS-212)",
        "Khadija Sajid (2k22-BSCS-256)",
    ]:
        add_center(name, 12, False)
    add_center("Supervised By", 12, False)
    add_center("Sir Syed Umair Asghar", 12, True)
    add_center("Head of Department", 12, False)
    add_center("Dr. Naeem Aslam", 12, True)
    add_center("DEPARTMENT OF COMPUTER SCIENCE", 12, True)
    add_center("NFC Institute of Engineering & Technology, Multan", 12, False)
    page_break()

    add_center("ACKNOWLEDGEMENT", 16, True)
    for kind, text in clean_paragraphs(pages_text.get(2, "")):
        if kind == "body" and "ACKNOWLEDGEMENT" not in text:
            add_just(text)
    page_break()

    add_center("ABSTRACT", 16, True)
    for kind, text in clean_paragraphs(pages_text.get(3, "")):
        if kind == "body" and text.upper() != "ABSTRACT":
            add_just(text)
    page_break()

    add_center("TABLE OF CONTENTS", 16, True)
    for level, title, page in TOC_ENTRIES:
        p = d.add_paragraph()
        p.paragraph_format.left_indent = DocxCm(0.5 * level)
        # tab leader
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab)
        pPr.append(tabs)
        run = p.add_run(title)
        run.font.name = "Times New Roman"
        run.font.size = DocxPt(12)
        run2 = p.add_run("\t" + page)
        run2.font.name = "Times New Roman"
        run2.font.size = DocxPt(12)
    page_break()

    add_center("LIST OF FIGURES", 16, True)
    for title, page in LOF:
        p = d.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab)
        pPr.append(tabs)
        r = p.add_run(title)
        r.font.name = "Times New Roman"
        r.font.size = DocxPt(12)
        r2 = p.add_run("\t" + page)
        r2.font.name = "Times New Roman"
        r2.font.size = DocxPt(12)
    page_break()

    add_center("LIST OF TABLES", 16, True)
    for title, page in LOT:
        p = d.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab)
        pPr.append(tabs)
        r = p.add_run(title)
        r.font.name = "Times New Roman"
        r.font.size = DocxPt(12)
        r2 = p.add_run("\t" + page)
        r2.font.name = "Times New Roman"
        r2.font.size = DocxPt(12)
    page_break()

    chapter_ranges = {
        1: range(11, 15),
        2: range(15, 29),
        3: range(29, 33),
        4: range(33, 39),
        5: range(39, 59),
        6: range(59, 61),
        7: range(61, 66),
        8: range(66, 70),
    }
    inserted = set()

    for page_no in range(11, 73):
        for num, name in CHAPTERS:
            if page_no == min(chapter_ranges[num]) and num not in inserted:
                add_center(f"CHAPTER {num}", 22, True)
                add_center(name, 18, True)
                page_break()
                inserted.add(num)
        if page_no == 70:
            add_center("CONCLUSION", 22, True)
            page_break()
        if page_no == 71:
            add_center("REFERENCES", 22, True)
            page_break()

        items = clean_paragraphs(pages_text.get(page_no, ""))
        if page_no in range(11, 15):
            items = renumber_chapter1(items)
        items = fix_unnumbered(items)

        i = 0
        while i < len(items):
            kind, text = items[i]
            if kind == "chapter":
                add_heading_left(text, 16)
            elif kind == "section":
                if text.startswith("2.6 Functional Requirements"):
                    add_heading_left(text, 14)
                    for b in FUNC_BULLETS:
                        p = d.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.left_indent = DocxCm(0.5)
                        run = p.add_run("• " + b)
                        run.font.name = "Times New Roman"
                        run.font.size = DocxPt(12)
                    i += 1
                    while i < len(items):
                        k2, t2 = items[i]
                        if k2 in {"section", "sub"} and (
                            t2.startswith("2.6.1") or t2.startswith("2.7")
                        ):
                            break
                        if k2 in {"section", "sub", "chapter", "caption"}:
                            break
                        i += 1
                    continue
                if text.startswith("2.7 Non Functional"):
                    add_heading_left(text, 14)
                    for b in NFR_BULLETS:
                        p = d.add_paragraph()
                        p.paragraph_format.left_indent = DocxCm(0.5)
                        run = p.add_run("• " + b)
                        run.font.name = "Times New Roman"
                        run.font.size = DocxPt(12)
                    i += 1
                    continue
                add_heading_left(text, 14)
            elif kind == "sub":
                if text.startswith("7.4"):
                    add_heading_left("7.4 Test Cases", 12)
                else:
                    add_heading_left(text, 12)
            elif kind == "caption":
                if text.startswith("Figure"):
                    img = largest_image_on_page(page_no)
                    if text.startswith("Figure 3.1"):
                        img = largest_image_on_page(30) or img
                    if img:
                        try:
                            d.add_picture(str(img), width=DocxInches(5.5))
                            d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception:
                            pass
                p = d.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.name = "Times New Roman"
                run.font.size = DocxPt(11)
                run.italic = True
            else:
                add_just(text)
            i += 1

    d.save(OUT_DOCX)
    print("Wrote DOCX", OUT_DOCX)


def main():
    ASSETS.mkdir(exist_ok=True)
    # ensure images exist
    if not any(ASSETS.glob("p*.png")):
        doc = pymupdf.open(SRC_PDF)
        for i, page in enumerate(doc):
            for j, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                try:
                    pix = pymupdf.Pixmap(doc, xref)
                    if pix.n > 4:
                        pix = pymupdf.Pixmap(pymupdf.csRGB, pix)
                    pix.save(str(ASSETS / f"p{i+1:02d}_{j}_{xref}.png"))
                except Exception:
                    pass
        doc.close()

    pages = load_pages_text()
    build_pdf(pages)
    build_docx_from_pdf_mirror()

    # quick sanity
    doc = pymupdf.open(OUT_PDF)
    print("PDF pages:", doc.page_count)
    # check no vertical-letter pages
    bad = 0
    for i in range(min(40, doc.page_count)):
        lines = [ln.strip() for ln in doc[i].get_text("text").splitlines() if ln.strip()]
        if lines and sum(1 for ln in lines if len(ln) == 1) > 20:
            bad += 1
            print("still-bad-page", i + 1)
    print("bad vertical pages in first 40:", bad)
    # sample chapter page
    for i in range(doc.page_count):
        t = doc[i].get_text("text").strip()
        if t.startswith("CHAPTER 1\n"):
            print("chapter opener @", i + 1, repr(t[:40]))
            break
    doc.close()


if __name__ == "__main__":
    main()
