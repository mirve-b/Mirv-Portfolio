#!/usr/bin/env python3
"""Apply FYP formatting to KAEL_FYP_Formatted.docx per plan."""

from __future__ import annotations

import copy
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Twips
from docx.text.paragraph import Paragraph

DOCX_PATH = "/Users/fatee/Downloads/KAEL_FYP_Formatted.docx"

FONT = "Times New Roman"
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_RIGHT = Cm(2.54)
MARGIN_LEFT = Cm(3.81)

CHAPTERS = [
    ("CHAPTER 1", "INTRODUCTION"),
    ("CHAPTER 2", "EXISTING SYSTEM"),
    ("CHAPTER 3", "SOFTWARE PROCESS MODEL"),
    ("CHAPTER 4", "PROPOSED SYSTEM"),
    ("CHAPTER 5", "SYSTEM DESIGN"),
    ("CHAPTER 6", "DEVELOPMENT"),
    ("CHAPTER 7", "SOFTWARE TESTING"),
    ("CHAPTER 8", "IMPLEMENTATION AND TRAINING"),
]

TOC_ENTRIES = [
    (0, "ACKNOWLEDGEMENT", "ii"),
    (0, "ABSTRACT", "iii"),
    (0, "TABLE OF CONTENTS", "v"),
    (0, "LIST OF FIGURES", "viii"),
    (0, "LIST OF TABLES", "x"),
    (0, "CHAPTER 1: INTRODUCTION", "1"),
    (1, "1.1 INTRODUCTION", "1"),
    (1, "1.2 Domain", "1"),
    (2, "Subdomain", "2"),
    (1, "1.3 Problem Statement", "2"),
    (1, "1.4 Motivation", "2"),
    (1, "1.5 Definition of Terms", "3"),
    (1, "1.6 Workflow", "3"),
    (1, "1.7 Goal of Project", "4"),
    (0, "CHAPTER 2: EXISTING SYSTEM", "5"),
    (1, "2.1 Relative Work", "5"),
    (1, "2.2 Existing Websites and their features", "5"),
    (1, "2.3 Novelty", "6"),
    (1, "2.4 Solution Available in KAEL", "6"),
    (1, "2.5 System Requirement Specification", "6"),
    (2, "2.5.1 User Interfaces", "6"),
    (2, "2.5.2 Hardware Interfaces", "6"),
    (2, "2.5.3 Software Interfaces", "7"),
    (1, "2.6 Functional Requirements", "7"),
    (2, "2.6.1 Use Case Diagram", "9"),
    (2, "2.6.2 Requirement Shells", "10"),
    (1, "2.7 Non Functional Requirements", "16"),
    (2, "2.7.1 Shell Requiremnets", "16"),
    (0, "CHAPTER 3: SOFTWARE PROCESS MODEL", "19"),
    (1, "3.1 Software Process Model", "19"),
    (1, "3.2 Proposed Model", "20"),
    (2, "3.2.1 Why Agile Scrum was chosen", "20"),
    (1, "3.3 Implementation of Agile Scrum", "21"),
    (0, "CHAPTER 4: PROPOSED SYSTEM", "23"),
    (1, "4.1 Need for the proposed system", "23"),
    (1, "4.2 Features of Proposed System", "23"),
    (2, "4.2.1 User Data Collection", "23"),
    (2, "4.2.2 AI Processing Pipeline", "24"),
    (2, "4.2.3 Automated generation of case studies", "24"),
    (2, "4.2.4 Portfolio Builder", "24"),
    (2, "4.2.5 ATS-Friendly CV Builder", "24"),
    (1, "4.3 Business Context", "24"),
    (2, "4.3.1 Advantages in Business", "25"),
    (2, "4.3.2 Stakeholder's Profiles", "26"),
    (2, "4.3.3 Project Priorities", "27"),
    (1, "4.4 Operating Environment", "27"),
    (2, "4.4.1 Client-Side Requirements", "27"),
    (2, "4.4.2 Server-Side / Backend Requirements", "28"),
    (2, "4.4.3 Security and Compliance Considerations", "28"),
    (0, "CHAPTER 5: SYSTEM DESIGN", "29"),
    (1, "5.1 Brief Description", "29"),
    (1, "5.2 UI/UX Design", "29"),
    (2, "1. Landing Page", "30"),
    (2, "2. Sign Up", "31"),
    (2, "3. Onboarding & Visual Identity Curation", "31"),
    (2, "4. Profile Creation", "33"),
    (2, "5. Case Study Generator", "35"),
    (2, "6. CV Generator", "36"),
    (2, "7. Portfolio Generator", "38"),
    (1, "5.3 ER Diagram for KAEL", "40"),
    (1, "5.4 DFD Diagram", "40"),
    (2, "Level 0", "40"),
    (2, "Level 1", "41"),
    (2, "Level 2", "42"),
    (1, "5.5 Activity Diagram KAEL", "43"),
    (1, "5.6 Class Diagram", "47"),
    (1, "5.7 Software Architectural Design", "47"),
    (1, "5.8 Sequence Diagram", "48"),
    (0, "CHAPTER 6: DEVELOPMENT", "49"),
    (1, "6.1 Selection of Tools & Technology", "49"),
    (2, "6.1.1 Hardware", "49"),
    (2, "6.1.2 Software", "49"),
    (0, "CHAPTER 7: SOFTWARE TESTING", "51"),
    (1, "7.1 Software Testing", "51"),
    (1, "7.2 Testing Process", "51"),
    (1, "7.3 Black Box Testing", "52"),
    (1, "7.4 Test Cases", "53"),
    (0, "CHAPTER 8: IMPLEMENTATION AND TRAINING", "56"),
    (1, "8.1 Implementation", "56"),
    (1, "8.2 Project Overview (Key Features only)", "57"),
    (2, "8.2.1 Expected Outcome", "57"),
    (1, "8.3 Training", "58"),
    (2, "8.3.1 Required Skillset", "58"),
    (2, "8.3.2 User Training Protocol", "58"),
    (0, "CONCLUSION", "60"),
    (0, "REFERENCES", "61"),
    (0, "FINAL YEAR PROJECT UNDERTAKING FORM", "62"),
]

LOF_ENTRIES = [
    ("Figure 1.1 KAEL Gantt Chart", "4"),
    ("Figure 2.1 Use case diagram - Student", "9"),
    ("Figure 2.2 Use case diagram - Recruiter", "9"),
    ("Figure 2.3 Use case diagram - Academic Authority", "10"),
    ("Figure 3.1 Agile scrum cycle adopted for KAEL development process", "20"),
    ("Figure 5.1 KAEL landing page", "30"),
    ("Figure 5.2 Sign-Up Screen", "31"),
    ("Figure 5.3 Creative-Academic Duality Slider", "32"),
    ("Figure 5.4 Profile Photo Upload", "32"),
    ("Figure 5.5 Educational Background Form", "32"),
    ("Figure 5.6 Soft Skills & Visual Identity", "33"),
    ("Figure 5.7 Soft Skills Selection", "33"),
    ("Figure 5.8 Profile Editor (Dark Theme)", "34"),
    ("Figure 5.9 Profile Editor (Light Theme)", "34"),
    ("Figure 5.10 Case Study Creation Interface", "35"),
    ("Figure 5.11 Generated Case Study Preview", "36"),
    ("Figure 5.12 CV Builder Form", "36"),
    ("Figure 5.13 Job Description Input", "37"),
    ("Figure 5.14 Job Match Analysis", "37"),
    ("Figure 5.15 Final ATS CV Preview", "37"),
    ("Figure 5.16 Portfolio Template Selection", "38"),
    ("Figure 5.17 Portfolio Live Preview", "39"),
    ("Figure 5.18 Portfolio Colour Customization", "39"),
    ("Figure 5.19 Customized Portfolio", "39"),
    ("Figure 5.20 ER Diagram of KAEL", "40"),
    ("Figure 5.21 DFD Level 0 - Context Diagram", "40"),
    ("Figure 5.22 DFD Level 1 - Onboarding, Case Study, Portfolio & Job Match Processes", "41"),
    ("Figure 5.23 DFD Level 2 - Detailed Process Decomposition", "42"),
    ("Figure 5.24 User Onboarding (Activity Diagram)", "43"),
    ("Figure 5.25 Case Study Creation (Activity Diagram)", "44"),
    ("Figure 5.26 Portfolio Generation & Publishing", "45"),
    ("Figure 5.27 Job Match & Tailored CV Generation", "46"),
    ("Figure 5.28 Class Diagram of KAEL", "47"),
    ("Figure 5.29 Software Architectural Design", "47"),
    ("Figure 5.30 Sequence Diagram of KAEL", "48"),
]

LOT_ENTRIES = [
    ("Table 1.1 Explanation of terms used in documentation", "3"),
    ("Table 2.1 Comparison with existing related web applications", "5"),
    ("Table 2.2 Software interfaces used", "7"),
    ("Table 4.1 Stakeholder's profiles", "26"),
    ("Table 4.2 Project Priorities of KAEL", "27"),
    ("Table 7.1 Test cases for KAEL", "53"),
]


def set_run_font(run, size=None, bold=None):
    run.font.name = FONT
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph: Paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, *, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def insert_paragraph_before(paragraph: Paragraph, text: str = "", **kwargs) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text or kwargs:
        set_paragraph_text(new_para, text, **kwargs)
    return new_para


def insert_paragraph_after(paragraph: Paragraph, text: str = "", **kwargs) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text or kwargs:
        set_paragraph_text(new_para, text, **kwargs)
    return new_para


def add_page_break(paragraph: Paragraph):
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def is_page_number_only(text: str) -> bool:
    t = text.strip()
    return bool(re.fullmatch(r"\d+|i{1,3}|iv|v|vi{0,3}|ix|x|xi{0,2}", t, re.I))


def apply_global_fonts_and_justify(doc: Document):
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.right_margin = MARGIN_RIGHT
        section.left_margin = MARGIN_LEFT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    for p in doc.paragraphs:
        text = p.text.strip()
        # Skip empty
        for run in p.runs:
            set_run_font(run, size=run.font.size.pt if run.font.size else 12)

        if not text:
            continue

        # Headings: chapter / numbered sections / list titles
        upper = text.upper()
        is_chapter = text.startswith("CHAPTER ") or upper in {
            "ACKNOWLEDGEMENT",
            "ABSTRACT",
            "TABLE OF CONTENTS",
            "LIST OF FIGURES",
            "LIST OF TABLES",
            "CONCLUSION",
            "REFERENCES",
        }
        is_section = bool(re.match(r"^\d+(\.\d+)*\s+\S", text)) or text in {
            "Subdomain",
            "INTRODUCTION",
        }

        if is_chapter and "\n" not in text and len(text) < 80:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, size=16, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        elif is_section and len(text) < 120 and not text.endswith("."):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, size=14, bold=True)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
        elif text.startswith("Figure ") or text.startswith("Table "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=11, bold=False)
                run.italic = True
        elif is_page_number_only(text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=12)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                if run.font.size is None:
                    set_run_font(run, size=12)
                else:
                    set_run_font(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        set_run_font(run, size=11)


def rebuild_leader_list(doc: Document, title_idx: int, entries, replace_until_pred):
    """Replace paragraphs after a list title with clean leader entries."""
    title_p = doc.paragraphs[title_idx]
    set_paragraph_text(title_p, title_p.text.strip().split("\n")[0], size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Delete following content until predicate
    to_delete = []
    for i in range(title_idx + 1, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if replace_until_pred(t, i):
            break
        to_delete.append(doc.paragraphs[i])
    for p in to_delete:
        p._element.getparent().remove(p._element)

    # Re-find title after deletions
    title_p = None
    for p in doc.paragraphs:
        if p.text.strip().upper() in {"LIST OF FIGURES", "LIST OF TABLES", "TABLE OF CONTENTS"}:
            # pick matching
            pass
    # Insert entries after title by searching again
    for p in doc.paragraphs:
        if p.text.strip().upper() == entries_title_name(entries, title_idx, doc):
            title_p = p
            break
    # Safer: use known titles
    return


def entries_title_name(entries, title_idx, doc):
    return doc.paragraphs[title_idx].text.strip().upper() if title_idx < len(doc.paragraphs) else ""


def find_para_by_exact(doc: Document, text: str):
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "\t" in t:
            continue
        if t == text or t.upper() == text.upper():
            return i, p
    return None, None


def find_para_startswith(doc: Document, prefix: str):
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "\t" in t:
            continue
        if t.startswith(prefix):
            return i, p
    return None, None


def add_leader_entry(after_para: Paragraph, title: str, page: str, indent_cm: float = 0) -> Paragraph:
    p = insert_paragraph_after(after_para, "")
    clear_paragraph(p)
    # Use tab with leader
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(int(Cm(14.0).twips)))  # within content width
    tabs.append(tab)
    pPr.append(tabs)

    run = p.add_run(title)
    set_run_font(run, size=12, bold=False)
    run2 = p.add_run("\t" + page)
    set_run_font(run2, size=12, bold=False)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent_cm)
    return p


def replace_block_with_leader_list(doc: Document, title_text: str, entries, stop_titles):
    idx, title_p = find_para_by_exact(doc, title_text)
    if title_p is None:
        idx, title_p = find_para_startswith(doc, title_text)
    if title_p is None:
        print(f"WARN: title not found: {title_text}")
        return

    # If title paragraph is a blob (title + entries), wipe it to title only first
    set_paragraph_text(title_p, title_text, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # collect deletions
    delete_elems = []
    collecting = False
    for p in doc.paragraphs:
        if p._p is title_p._p:
            collecting = True
            continue
        if not collecting:
            continue
        t = p.text.strip()
        if any(t.startswith(s) or t.upper() == s.upper() for s in stop_titles):
            break
        delete_elems.append(p._element)

    for el in delete_elems:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # re-find title
    _, title_p = find_para_by_exact(doc, title_text)
    cursor = title_p
    for item in entries:
        if len(item) == 3:
            level, name, page = item
            indent = {0: 0, 1: 0.5, 2: 1.0}.get(level, 0)
        else:
            name, page = item
            indent = 0
        cursor = add_leader_entry(cursor, name, page, indent_cm=indent)
    print(f"Rebuilt {title_text} with {len(entries)} entries")


def insert_chapter_title_pages(doc: Document):
    # Process from last to first so indices remain valid for earlier chapters
    for num, name in reversed(CHAPTERS):
        needle = f"{num}: {name}"
        target = None
        for p in doc.paragraphs:
            t = p.text.strip()
            # Skip TOC leader lines (contain tab + page number)
            if "\t" in t:
                continue
            # Match body chapter headings only
            if t == needle or t.startswith(needle + " ") or t.startswith(needle + "\n"):
                target = p
                break
            # Merged heading+body on one line
            if t.startswith(needle) and len(t) > len(needle):
                target = p
                break
        if target is None:
            # Fallback: CHAPTER N: at start, no tab
            for p in doc.paragraphs:
                t = p.text.strip()
                if "\t" in t:
                    continue
                if t.startswith(f"{num}:") and name[:10].upper() in t.upper():
                    target = p
                    break
        if target is None:
            print(f"WARN: chapter not found {needle}")
            continue

        # If paragraph merges chapter + body, split heading out
        t = target.text.strip()
        if "\n" in t or len(t) > len(needle) + 5:
            # Keep only chapter line as heading, move rest to following para
            parts = re.split(r"\n+", t, maxsplit=1)
            heading = parts[0].strip()
            rest = parts[1].strip() if len(parts) > 1 else ""
            # Also handle "CHAPTER 5: SYSTEM DESIGN 5.1 Brief..." without newline
            m = re.match(r"(CHAPTER\s+\d+:\s+[A-Z][A-Z\s/&-]+?)(\s+\d+\.\d+.*)$", heading)
            if m:
                heading, rest2 = m.group(1).strip(), m.group(2).strip()
                rest = (rest2 + ("\n" + rest if rest else "")).strip()
            set_paragraph_text(target, heading, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            if rest:
                insert_paragraph_after(target, rest, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # Insert title page BEFORE chapter heading: page break + centered content + page break
        # Structure: [page break][CHAPTER N][NAME][page break][original heading]
        spacer = insert_paragraph_before(target, "")
        add_page_break(spacer)

        # Vertical centering approximation: several empty paragraphs
        cursor = spacer
        for _ in range(8):
            cursor = insert_paragraph_after(cursor, "")

        ch_line = insert_paragraph_after(cursor, num, size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        name_line = insert_paragraph_after(ch_line, name, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        after = insert_paragraph_after(name_line, "")
        for _ in range(8):
            after = insert_paragraph_after(after, "")
        add_page_break(after)
        print(f"Inserted title page for {num}: {name}")


def fix_heading_numbers_and_splits(doc: Document):
    replacements = [
        ("Server-Side / Backend Requirements", "4.4.2 Server-Side / Backend Requirements"),
        ("Security and Compliance Considerations", "4.4.3 Security and Compliance Considerations"),
        ("Expected Outcome", "8.2.1 Expected Outcome"),
        ("Required Skillset", "8.3.1 Required Skillset"),
        ("User Training Protocol", "8.3.2 User Training Protocol"),
    ]

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "\t" in t:
            continue
        # 7.4 split heading
        if re.match(r"7\.4\s*Test\s*\n?\s*Cases", t) or (t.startswith("7.4 Test") and "Cases" in t):
            body = ""
            m = re.search(r"7\.4\s*Test\s*Cases\s*(.*)$", t, re.S)
            if m:
                body = m.group(1).strip()
            if t == "7.4 Test Cases" and not body:
                continue
            set_paragraph_text(p, "7.4 Test Cases", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            if body:
                insert_paragraph_after(p, body, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            print("Fixed 7.4 Test Cases heading")

        # Number unnumbered headings (only when they start the paragraph)
        for old, new in replacements:
            if t == old or t.startswith(old + " ") or t.startswith(old + "\n"):
                rest = t[len(old) :].lstrip(" \n")
                if rest:
                    set_paragraph_text(p, new, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
                    insert_paragraph_after(p, rest, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                else:
                    set_paragraph_text(p, new, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
                print(f"Numbered heading: {new}")
                break

    # Insert Client-Side heading before client content block if missing in body
    has_client = any(
        p.text.strip().startswith("4.4.1 Client-Side") and "\t" not in p.text
        for p in doc.paragraphs
    )
    if not has_client:
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.startswith("KAEL is a desktop app created with Flutter") or (
                t.startswith("The application must have an active internet connection")
            ):
                insert_paragraph_before(
                    p,
                    "4.4.1 Client-Side Requirements",
                    size=14,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                )
                print("Inserted 4.4.1 Client-Side Requirements")
                break


def bulletize_functional_requirements(doc: Document):
    """Add bullets to 2.6 feature list items."""
    in_block = False
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if "\t" in t:
            continue
        if t.startswith("2.6 Functional Requirements"):
            in_block = True
            continue
        if in_block and (t.startswith("2.6.1") or t.startswith("2.7")):
            break
        if not in_block:
            continue
        if not t or is_page_number_only(t):
            continue
        # Feature lines like "Profile Management: ..."
        if ":" in t and not t.startswith("Requirement") and not t.startswith("•"):
            clear_paragraph(p)
            run = p.add_run("• " + t)
            set_run_font(run, size=12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(6)
            count += 1
    print(f"Bulleted {count} functional requirement list items")


def fix_nonfunctional_bullets(doc: Document):
    """Ensure 2.7 intro list uses small bullets."""
    idx, p = find_para_startswith(doc, "2.7 Non Functional")
    if p is None:
        return
    # Insert small-bullet category list after heading if not present
    nxt = None
    # find next para
    found = False
    for i, q in enumerate(doc.paragraphs):
        if q._p is p._p:
            found = True
            continue
        if found:
            nxt = q
            break
    categories = [
        "Security & Data Protection",
        "Performance & Responsiveness",
        "Scalability & Availability",
        "Usability & Accessibility",
    ]
    # If next content is 2.7.1, insert categories
    if nxt and nxt.text.strip().startswith("2.7.1"):
        cursor = p
        for cat in categories:
            cursor = insert_paragraph_after(cursor, "• " + cat, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            cursor.paragraph_format.left_indent = Cm(0.5)
            cursor.paragraph_format.space_after = Pt(2)
        print("Inserted small-bullet NFR categories under 2.7")
    else:
        # Replace any oversized bullet characters in nearby paragraphs
        for q in doc.paragraphs:
            if "●" in q.text:
                new_t = q.text.replace("●", "•")
                set_paragraph_text(q, new_t.strip(), size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
                q.paragraph_format.left_indent = Cm(0.5)
        print("Normalized oversized bullet characters in NFR section")


def remove_extra_blank_spaces(doc: Document):
    """Collapse runs of empty paragraphs (keep at most one)."""
    empties = []
    streak = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        # also treat lone page numbers as removable noise in body (keep centered footer-like sparingly)
        if not t:
            streak += 1
            if streak > 1:
                empties.append(p)
        else:
            streak = 0

    for p in empties:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
    print(f"Removed {len(empties)} extra blank paragraphs")


def move_figure_31_up(doc: Document):
    """Place Figure 3.1 closer under 3.2 Proposed Model to fill sparse space."""
    _, proposed = find_para_startswith(doc, "3.2 Proposed Model")
    fig_idx, fig = find_para_startswith(doc, "Figure 3.1")
    if not proposed or not fig:
        print("WARN: could not relocate Figure 3.1")
        return
    # If figure already immediately follows proposed model (within a few paras), skip
    # Move figure element to after proposed model
    fig_el = fig._element
    # Also move following image drawing paragraph if empty-with-drawing after figure caption is before
    # Find drawing paragraphs near figure: previous or next with drawings
    neighbors = []
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if p._p is fig._p:
            # previous empty/drawing paras
            j = i - 1
            while j >= 0 and (not paras[j].text.strip() or paras[j]._p.xpath(".//w:drawing")):
                if paras[j]._p.xpath(".//w:drawing"):
                    neighbors.insert(0, paras[j]._p)
                j -= 1
            # next drawing-only
            j = i + 1
            while j < len(paras) and (not paras[j].text.strip() or paras[j]._p.xpath(".//w:drawing")):
                if paras[j]._p.xpath(".//w:drawing"):
                    neighbors.append(paras[j]._p)
                else:
                    break
                j += 1
            break

    # Detach and reattach after proposed
    anchor = proposed._p
    for el in neighbors:
        if el.getparent() is not None:
            el.getparent().remove(el)
            anchor.addnext(el)
            anchor = el
    if fig_el.getparent() is not None:
        fig_el.getparent().remove(fig_el)
        anchor.addnext(fig_el)
    print("Moved Figure 3.1 under 3.2 Proposed Model")


def remove_blank_page_iv(doc: Document):
    """Remove lone 'iv' blank page marker paragraph if present."""
    for p in list(doc.paragraphs):
        if p.text.strip() == "iv":
            # remove surrounding empties too
            p._element.getparent().remove(p._element)
            print("Removed blank page iv marker")
            break


def main():
    print("Loading", DOCX_PATH)
    doc = Document(DOCX_PATH)

    remove_blank_page_iv(doc)
    apply_global_fonts_and_justify(doc)

    # Rebuild TOC / LOF / LOT
    replace_block_with_leader_list(
        doc,
        "TABLE OF CONTENTS",
        TOC_ENTRIES,
        stop_titles=["LIST OF FIGURES", "CHAPTER 1"],
    )
    # TOC title may have been a blob — handle blob start
    # If TABLE OF CONTENTS not exact, find blob
    if not any(p.text.strip() == "TABLE OF CONTENTS" for p in doc.paragraphs):
        for p in doc.paragraphs:
            if p.text.strip().startswith("TABLE OF CONTENTS"):
                set_paragraph_text(p, "TABLE OF CONTENTS", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                break
        replace_block_with_leader_list(
            doc,
            "TABLE OF CONTENTS",
            TOC_ENTRIES,
            stop_titles=["LIST OF FIGURES", "CHAPTER 1"],
        )

    replace_block_with_leader_list(
        doc,
        "LIST OF FIGURES",
        LOF_ENTRIES,
        stop_titles=["LIST OF TABLES", "CHAPTER 1"],
    )
    replace_block_with_leader_list(
        doc,
        "LIST OF TABLES",
        LOT_ENTRIES,
        stop_titles=["CHAPTER 1", "CHAPTER 1: INTRODUCTION"],
    )

    fix_heading_numbers_and_splits(doc)
    bulletize_functional_requirements(doc)
    fix_nonfunctional_bullets(doc)
    move_figure_31_up(doc)
    remove_extra_blank_spaces(doc)
    insert_chapter_title_pages(doc)

    # Re-apply fonts after insertions
    apply_global_fonts_and_justify(doc)

    doc.save(DOCX_PATH)
    print("Saved", DOCX_PATH)


if __name__ == "__main__":
    main()
